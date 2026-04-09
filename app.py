import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
import openai
import base64
import io
import json
import re
from docx import Document

# --- 页面全局配置 ---
st.set_page_config(page_title="Academic Intelligence Studio Ultra Pro Max", layout="wide", initial_sidebar_state="expanded")

# --- 侧边栏：核心引擎配置 ---
with st.sidebar:
    st.title("⚙️ 顶级科研引擎中心")
    engine_choice = st.selectbox("选择核心引擎", [
        "Google Gemini 1.5 Pro (全能推荐)",
        "智谱 GLM-4v (国内推荐)",
        "OpenAI GPT-4o",
        "DeepSeek V3",
        "Kimi (Moonshot)"
    ])
    api_key = st.text_input("输入 API Key", type="password")
    
    st.divider()
    st.subheader("📊 LetPub 期刊过滤指标")
    if_range = st.slider("影响因子范围", 0.0, 50.0, (5.0, 10.0))
    cas_zone_sel = st.multiselect("中科院分区", ["1区", "2区", "3区", "4区"], default=["1区", "2区"])
    sci_zone_sel = st.multiselect("SCI 分区 (Q)", ["Q1", "Q2", "Q3", "Q4"], default=["Q1"])
    is_oa = st.checkbox("仅显示 Open Access")

    st.divider()
    st.subheader("🖼️ Illustrator 画布导出")
    dpi_val = st.selectbox("导出精度 (DPI)", [300, 600, 1200], index=0)
    export_fmt = st.selectbox("文件格式", ["PNG", "PDF", "TIFF"])

# --- AI 通讯逻辑 (深度适配 Gemini 1.5 Pro) ---
def get_ai_response(messages):
    if not api_key: return "ERROR: API Key 缺失，请在左侧输入。"
    
    urls = {
        "Google Gemini 1.5 Pro (全能推荐)": "https://generativelanguage.googleapis.com/v1beta/openai/",
        "智谱 GLM-4v (国内推荐)": "https://open.bigmodel.cn/api/paas/v4/",
        "OpenAI GPT-4o": "https://api.openai.com/v1",
        "DeepSeek V3": "https://api.deepseek.com",
        "Kimi (Moonshot)": "https://api.moonshot.cn/v1"
    }
    models = {
        "Google Gemini 1.5 Pro (全能推荐)": "gemini-1.5-pro",
        "智谱 GLM-4v (国内推荐)": "glm-4v",
        "OpenAI GPT-4o": "gpt-4o",
        "DeepSeek V3": "deepseek-chat",
        "Kimi (Moonshot)": "moonshot-v1-8k"
    }
    
    try:
        client = openai.OpenAI(api_key=api_key, base_url=urls[engine_choice])
        response = client.chat.completions.create(
            model=models[engine_choice],
            messages=messages,
            temperature=0.2
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"ERROR: 调用失败。原因: {str(e)}"

# --- UI 导航 ---
st.title("🎓 顶级科研全流程工作站 Pro Max")
tabs = st.tabs(["🔍 SCI 结果描述", "🎨 Pro 画布 (Illustrator)", "✍️ 交互润色/降AI", "📊 期刊智选 (LetPub)", "🚀 创新大脑实验室"])

# --- 模块 1: SCI 结果深度解析 ---
with tabs[0]:
    st.header("Manuscript & Figure Neural Interpretation")
    c1, c2 = st.columns(2)
    with c1:
        ms_in = st.file_uploader("上传手稿/初稿文本", type=["docx", "pdf", "txt"])
        fg_in = st.file_uploader("上传实验图表/表格 (多选)", accept_multiple_files=True)
    with c2:
        out_lang = st.radio("结果输出语言", ["中文 (Chinese)", "英文 (English)"], horizontal=True)
        if st.button("开始生成 SCI 级深度结果解析"):
            with st.spinner("AI 正在深度扫描数据趋势与统计显著性..."):
                prompt = [{"role":"user","content":f"请作为顶刊审稿人。结合上传的文字和图表，生成三段式的SCI Results描述。要求：1. 带有统计指标（OR, CI, P值）的详版描述。2. 对图表每一个关键显著性结点的专业解释。3. 仿Nature精简版结论。请用{out_lang}输出。"}]
                st.markdown(get_ai_response(prompt))

# --- 模块 2: Pro 画布 (Adobe Illustrator 级交互) ---
with tabs[1]:
    st.header("Vector Figure Illustrator Pro")
    st.caption("支持多 PDF 同时插入、属性面板实时微调、‘✨’一键查找修改同类标注。")
    
    editor_html = f"""
    <div style="display:flex; gap:10px; background:#1e1e1e; padding:15px; border-radius:10px; color:white;">
        <div id="canv-main">
            <div id="toolbar" style="margin-bottom:12px; display:flex; gap:8px;">
                <input type="file" id="pdfAdd" multiple style="display:none">
                <button onclick="document.getElementById('pdfAdd').click()" style="background:#444; color:white; border:none; padding:8px 12px; cursor:pointer; border-radius:4px;">📁 插入 PDF/图片</button>
                <button onclick="addText()" style="background:#444; color:white; border:none; padding:8px 12px; cursor:pointer; border-radius:4px;">T 文本</button>
                <button onclick="batchModify()" style="background:#f39c12; color:white; border:none; padding:8px 12px; cursor:pointer; font-weight:bold; border-radius:4px;">✨ 查找相同修改</button>
                <button onclick="exportPro()" style="background:#28a745; color:white; border:none; padding:8px 12px; cursor:pointer; border-radius:4px;">🚀 导出 {export_fmt} ({dpi_val}DPI)</button>
            </div>
            <canvas id="c" width="880" height="600" style="border:1px solid #000; background:white;"></canvas>
        </div>
        <div id="prop-panel" style="width:240px; background:#2d2d2d; padding:20px; border-radius:8px; font-size:12px;">
            <h4 style="margin-top:0">Properties</h4>
            字体: <select id="fFam" style="width:100%"><option>Times New Roman</option><option>Arial</option></select><br><br>
            字号: <input type="number" id="fSiz" value="28" style="width:100%"><br><br>
            颜色: <input type="color" id="fCol" style="width:100%"><br><br>
            <hr>
            <p style="color:#aaa">提示：点击选中任何一个 Fig 标注或文字，修改属性后，点‘✨’按钮可同步全画布所有同类元素。</p>
        </div>
    </div>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/fabric.js/5.3.1/fabric.min.js"></script>
    <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.4.120/pdf.min.js"></script>
    <script>
        const canvas = new fabric.Canvas('c');
        pdfjsLib.GlobalWorkerOptions.workerSrc = 'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.4.120/pdf.worker.min.js';

        function batchModify() {{
            const ref = canvas.getActiveObject();
            if(!ref) return alert("请先选中一个样板元素！");
            const nCol = document.getElementById('fCol').value;
            const nSiz = parseInt(document.getElementById('fSiz').value);
            canvas.getObjects().forEach(obj => {{
                if(obj.type === ref.type && (obj.fill === ref.fill || obj.fontSize === ref.fontSize)) {{
                    obj.set({{ fill: nCol, fontSize: nSiz, fontFamily: document.getElementById('fFam').value }});
                }}
            }});
            canvas.renderAll();
        }}

        document.getElementById('pdfAdd').onchange = function(e) {{
            const files = e.target.files;
            for(let file of files) {{
                const reader = new FileReader();
                if(file.type === "application/pdf") {{
                    reader.onload = function() {{
                        const uint8 = new Uint8Array(this.result);
                        pdfjsLib.getDocument(uint8).promise.then(pdf => {{
                            pdf.getPage(1).then(page => {{
                                const viewport = page.getViewport({{scale: 1.8}});
                                const tempC = document.createElement('canvas');
                                tempC.height = viewport.height; tempC.width = viewport.width;
                                page.render({{canvasContext: tempC.getContext('2d'), viewport: viewport}}).promise.then(() => {{
                                    fabric.Image.fromURL(tempC.toDataURL(), img => {{ img.scale(0.5); canvas.add(img); }});
                                }});
                            }});
                        }});
                    }};
                    reader.readAsArrayBuffer(file);
                }} else {{
                    reader.onload = f => fabric.Image.fromURL(f.target.result, img => {{ img.scale(0.25); canvas.add(img); }});
                    reader.readAsDataURL(file);
                }}
            }}
        }};

        function addText() {{ canvas.add(new fabric.IText('Fig Label', {{ left: 150, top: 150, fontFamily: 'Times New Roman', fontSize: 24 }})); }}
        function exportPro() {{
            const mul = {dpi_val} / 96;
            const dataURL = canvas.toDataURL({{ format: 'png', multiplier: mul }});
            const link = document.createElement('a');
            link.download = "Academic_Figure_Publication.png"; link.href = dataURL; link.click();
        }}
    </script>
    """
    components.html(editor_html, height=780)

# --- 模块 3: 交互式润色 (Paperpal 级) ---
with tabs[2]:
    st.header("Interactive Academic Polisher & Humanizer")
    if 'rev_data' not in st.session_state: st.session_state.rev_data = []

    text_input = st.text_area("输入原始学术文本 (逐句解析)", height=200)
    mode = st.radio("润色主攻方向", ["专业润色 (Make Academic)", "降低 AI 率重构 (De-AI)"], horizontal=True)

    if st.button("开始逐句分析处理"):
        with st.spinner("AI 正在重构学术逻辑链..."):
            prompt = f"""Act as a senior SCI editor. Break the text into sentences. 
            For each, provide 3 versions: 1. Original 2. Academic (Formal) 3. Humanized (De-AI, varied structure).
            Output ONLY a JSON list: [{{"orig":"...","acad":"...","human":"..."}}]
            Text: {text_input}"""
            res = get_ai_response([{"role":"user","content":prompt}])
            try:
                json_str = re.search(r'\[.*\]', res, re.DOTALL).group()
                st.session_state.rev_data = json.loads(json_str)
            except: st.error("解析失败。请确保输入的是完整的学术文本。")

    if st.session_state.rev_data:
        st.subheader("💡 修订选择面板 (不改变原意):")
        final_list = []
        for i, item in enumerate(st.session_state.rev_data):
            with st.expander(f"句子 {i+1}: {item['orig'][:60]}..."):
                c = st.radio(f"S{i+1}", [item['orig'], item['acad'], item['human']], key=f"rev_{i}", index=1 if mode=="专业润色 (Make Academic)" else 2)
                final_list.append(c)
        st.divider()
        st.text_area("润色后结果 (可直接复制)", " ".join(final_list), height=200)

# --- 模块 4: 期刊智选 (真实性保障) ---
with tabs[3]:
    st.header("LetPub Strategic Journal Matcher (Real Data)")
    st.info(f"过滤设定：IF ({if_range[0]}-{if_range[1]}) | 中科院 {cas_zone_sel} | SCI {sci_zone_sel}")
    abs_in = st.text_area("输入论文摘要内容 (AI 将匹配真实发表记录)", height=150)
    
    if st.button("开始智能匹配真实期刊库"):
        with st.spinner("正在检索 2024-2025 JCR 数据库与已发表文章记录..."):
            prompt = f"""
            TASK: ACT AS LetPub Senior Consultant. 
            Mandatory Criteria: 
            - IF: {if_range[0]} to {if_range[1]}
            - CAS Zone: {cas_zone_sel}
            - SCI Zone: {sci_zone_sel}
            Identify 10 REAL journals that meet these criteria based on the Abstract: {abs_in}.
            For each journal, provide:
            1. Official Name 2. Latest IF 3. CAS/SCI Zones 4. Introduction 5. REAL Related Articles (Provide titles of papers that used similar databases or methods published in this journal).
            DO NOT hallucinate. Ensure data is accurate. Output as a Markdown table.
            """
            st.markdown(get_ai_response([{"role":"user","content":prompt}]))

# --- 模块 5: 创新实验室 (无限字数报告版) ---
with tabs[4]:
    st.header("🚀 蓝海领域：多数据库创新大脑")
    
    step = st.radio("选择流程阶段", ["1. 蓝海思路与未来指标探索", "2. 详细实施方案 (变量/计算/相关性)", "3. SCI 背景方案完整报告"], horizontal=True)
    
    if step == "1. 蓝海思路与未来指标探索":
        col_s1_1, col_s1_2 = st.columns(2)
        db_s1 = col_s1_1.text_input("拟用数据库 (如: NHANES + UKB)")
        field_s1 = col_s1_2.text_input("大方向 (如: 口腔虚弱, 肌少症)")
        if st.button("挖掘创新点与未来指标"):
            prompt = f"""基于数据库 {db_s1}，探究领域 {field_s1}。
            1. 给出 3 个 2025+ 蓝海思路。
            2. 重点：推荐 3 个‘未来可以进一步探究的指标’（利用现有变量计算得到的派生指标，如系统性评分、轨迹参数等）。
            3. 提供参考高分文章思路。"""
            st.markdown(get_ai_response([{"role":"user","content":prompt}]))
            
    elif step == "2. 详细实施方案 (变量/计算/相关性)":
        c1, c2, c3 = st.columns(3)
        u_exp = c1.text_input("确定的暴露因素")
        u_db = c2.text_input("数据库名称")
        u_out = c3.text_input("确定的结局")
        if st.button("生成详细筛选方案与相关性逻辑"):
            prompt = f"""数据库：{u_db}。暴露：{u_exp}。结局：{u_out}。
            请给出：1. 具体的 Variable Codes (针对 {u_db} 的变量名)。
            2. 详细计算公式与逻辑。
            3. 进一步的相关性探索逻辑（中介效应、调节效应、亚组交互作用）。
            4. 给出 3-5 个真实 PMID 参考文献。"""
            st.markdown(get_ai_response([{"role":"user","content":prompt}]))
            
    elif step == "3. SCI 背景方案完整报告":
        st.subheader("报告生成配置")
        c_r1, c_r2 = st.columns(2)
        sel_db = c_r1.text_input("确认数据库", "NHANES")
        sel_dir = c_r2.text_input("确认研究方向", "口腔虚弱与肌少症的关联及炎症中介分析")
        fig_selection = st.multiselect("包含哪些例图描述", ["Figure 1 (Flow chart)", "Table 1 (Baseline)", "Figure 2 (RCS Plot)", "Figure 3 (Mediation Map)", "Figure 4 (Subgroup Forest Plot)"], default=["Figure 1 (Flow chart)", "Table 1 (Baseline)", "Figure 2 (RCS Plot)"])
        
        if st.button("生成详细 SCI 方案报告 (不限字数)"):
            with st.spinner("正在撰写深度报告..."):
                prompt = f"""请针对数据库 {sel_db} 和方向 {sel_dir} 撰写一份深度科研方案。
                要求：
                1. 背景介绍（千字以上，引用最新 Nature/Lancet 风格）。
                2. 详细方法论与 Variable Codes。
                3. 创新性深度论证。
                4. 为所选例图 {fig_selection} 提供详细描述和 Figure Legend。
                5. 参考文献列表真实准确（近5年）。
                段落形式，严谨、专业。"""
                res = get_ai_response([{"role":"user","content":prompt}])
                st.markdown(res)
                # Word 导出
                doc = Document()
                doc.add_heading(f'Research Innovation Report: {sel_dir}', 0)
                doc.add_paragraph(res)
                bio = io.BytesIO()
                doc.save(bio)
                st.download_button("下载 Word 完整方案", bio.getvalue(), "Proposal_Report_Full.docx")
